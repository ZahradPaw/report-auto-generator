from PyQt6.QtWidgets import QMainWindow, QMessageBox, QFileDialog
from PyQt6.QtGui import QIcon
from datetime import datetime
import docx
from docx.shared import Pt
import pickle
import docx2pdf
import os
import sys
from ui_mainwindow import Ui_MainWindow

popular_russian_cities = ['Москва', 'Санкт-Петербург', 'Новосибирск', 'Екатеринбург', 'Казань', 'Нижний Новгород', 'Челябинск', 'Самара', 'Омск', 'Ростов-на-Дону', 'Уфа', 'Красноярск', 'Пермь', 'Воронеж', 'Волгоград', 'Краснодар', 'Саратов', 'Тюмень', 'Тольятти', 'Ижевск', 'Барнаул', 'Ульяновск', 'Иркутск', 'Хабаровск', 'Ярославль', 'Владивосток', 'Махачкала', 'Томск', 'Оренбург', 'Кемерово', 'Новокузнецк', 'Рязань', 'Астрахань', 'Набережные Челны', 'Пенза', 'Липецк', 'Киров', 'Чебоксары', 'Тула', 'Калининград', 'Балашиха', 'Курск', 'Севастополь', 'Сочи', 'Ставрополь', 'Улан-Удэ', 'Архангельск', 'Чита', 'Смоленск', 'Орёл']


class MainWindow(QMainWindow):
    """Класс главного окна"""

    def __init__(self):
        super(MainWindow, self).__init__()

        # Загрузка интерфейса
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        # pyuic6 -x mainwindow.ui -o ui_mainwindow.py

        self.icon = QIcon(self.resource_path('icon.ico'))
        self.setWindowIcon(self.icon)
        self.setGeometry(100, 100, 500, 500)

        self.ui.questionButton.setIcon(QIcon(self.resource_path('question_icon.png')))

        # Кол-во студентов в отчете
        self.students_count = 1
        self.students_labels = (self.ui.studentLabel2, self.ui.studentLabel3, self.ui.studentLabel4, self.ui.studentLabel5)
        self.students_edits = (self.ui.studentEdit2, self.ui.studentEdit3, self.ui.studentEdit4, self.ui.studentEdit5)

        # Содержание списков
        self.ui.workTypeBox.addItems(['Лабораторная', 'Практическая'])
        self.ui.cityBox.addItems(popular_russian_cities)

        # Функции кнопок переходов по страницам
        self.ui.nextButton1.clicked.connect(lambda: self.ui.tabWidget.setCurrentIndex(1))
        self.ui.nextButton2.clicked.connect(lambda: self.ui.tabWidget.setCurrentIndex(2))
        self.ui.nextButton3.clicked.connect(lambda: self.ui.tabWidget.setCurrentIndex(3))
        self.ui.backButton1.clicked.connect(lambda: self.ui.tabWidget.setCurrentIndex(0))
        self.ui.backButton2.clicked.connect(lambda: self.ui.tabWidget.setCurrentIndex(1))
        self.ui.backButton3.clicked.connect(lambda: self.ui.tabWidget.setCurrentIndex(2))

        # Функции кнопок +/-
        self.ui.minusButton.clicked.connect(self.change_students_count)
        self.ui.plusButton.clicked.connect(self.change_students_count)

        # Функция кнопки с информацией о шаблоне
        self.ui.questionButton.clicked.connect(self.show_template_info)

        # Путь к файлу-шаблону отчета
        self.template_file_url = ''
        self.report_file_name = ''
        self.report_format = ''

        # Функция кнопки изменения пути к шаблону отчета
        self.ui.changeButton.clicked.connect(self.change_template_file)

        # Функция кнопки создания отчета
        self.ui.createReportButton.clicked.connect(self.create_report)

        # Список форматов отчета
        self.ui.formatBox.addItems(['.docx - Word документ', '.doc - Word документ (1997 - 2004)', '.pdf - электронный документ'])

        # Список данных для отчета
        self.templates = {
                     '{university_name}': '',
                     '{faculty_name}': '',
                     '{work_type}': '',
                     '{work_num}': '1',
                     '{subject_name}': '',
                     '{work_topic}': '',
                     '{group_num}': '',
                     '{student_name1}': '-',
                     '{student_name2}': '-',
                     '{student_name3}': '-',
                     '{student_name4}': '-',
                     '{student_name5}': '-',
                     '{teacher_post}': '',
                     '{teacher_name}': '',
                     '{city}': '',
                     '{year}': '',
                     '{purpose_work}': '',
                     '{work_progress}': '',
                     '{conclusion}': ''
        }

        # Загрузка сохраненных данных
        self.load_saved_data()

        # Изменение кол-ва студентов в отчете
        self.change_students_count()

    def load_saved_data(self):
        """Загрузка сохраненных данных из файла"""
        try:
            with open('data.bin', 'rb') as data:
                self.templates = pickle.load(data)
                [self.students_count, self.template_file_url, self.report_format, self.report_file_name] = pickle.load(data)
                # Загрузка значений в поля
                self.ui.universtyEdit.setText(self.templates['{university_name}'])
                self.ui.facultyEdit.setText(self.templates['{faculty_name}'])
                self.ui.workTypeBox.setCurrentIndex(0 if self.templates['{work_type}'] == 'лабораторной' else 1)
                self.ui.workNumberBox.setValue(int(self.templates['{work_num}']))
                self.ui.subjectEdit.setText(self.templates['{subject_name}'])
                self.ui.topicEdit.setText(self.templates['{work_topic}'])
                self.ui.groupEdit.setText(self.templates['{group_num}'])
                self.ui.studentEdit1.setText(self.templates['{student_name1}'].strip('\n'))
                self.ui.studentEdit2.setText(self.templates['{student_name2}'].strip('\n'))
                self.ui.studentEdit3.setText(self.templates['{student_name3}'].strip('\n'))
                self.ui.studentEdit4.setText(self.templates['{student_name4}'].strip('\n'))
                self.ui.studentEdit5.setText(self.templates['{student_name5}'].strip('\n'))
                self.ui.studentCountLabel.setText(str(self.students_count))
                self.ui.postEdit.setText(self.templates['{teacher_post}'])
                self.ui.teacherEdit.setText(self.templates['{teacher_name}'])
                self.ui.cityBox.setCurrentText(self.templates['{city}'])
                self.ui.purposeText.setText(self.templates['{purpose_work}'])
                self.ui.progressText.setText(self.templates['{work_progress}'])
                self.ui.conclusionText.setText(self.templates['{conclusion}'])
                self.ui.templateFileEdit.setText(self.template_file_url)
                self.ui.fileNameEdit.setText(self.report_file_name)
                self.ui.formatBox.setCurrentText(self.report_format)
        except Exception as e:
            print(f'Ошибка при загрузке данных: {str(e)}')

    def save_data(self):
        """Сохранение данных полей в файл"""
        try:
            with open('data.bin', 'wb') as data:
                pickle.dump(self.templates, data)
                pickle.dump([self.students_count, self.template_file_url, self.report_format, self.report_file_name], data)
        except Exception as e:
            print(f'Ошибка при сохранении данных: {str(e)}')

    def change_students_count(self):
        """Изменение кол-ва студентов в отчете"""
        action = self.sender()
        if action is not None:
            if action.text() == '-':
                self.students_count = self.students_count - 1 if self.students_count > 1 else self.students_count
            elif action.text() == '+':
                self.students_count = self.students_count + 1 if self.students_count < 5 else self.students_count
        self.ui.studentCountLabel.setText(str(self.students_count))
        # Блокировка виджетов студентов
        for i in range(3, self.students_count - 2, -1):
            self.students_labels[i].setStyleSheet('color: gray;')
            self.students_edits[i].setStyleSheet('color: gray;')
            self.students_edits[i].setReadOnly(True)
        # Разблокировка виджетов студентов
        for i in range(0, self.students_count - 1):
            self.students_labels[i].setStyleSheet('color: #1A3E72;')
            self.students_edits[i].setStyleSheet('color: black;')
            self.students_edits[i].setReadOnly(False)

    def change_template_file(self):
        """Выбор нового файла шаблона для отчета"""
        pattern_file = QFileDialog.getOpenFileName(self, 'Выбор файла-шаблона отчета',
                                                   f"{os.environ['USERPROFILE']}\\Documents", 'Word (*.docx *.doc)')[0]

        if self.is_valid_pattern(pattern_file):
            self.template_file_url = pattern_file
            self.ui.templateFileEdit.setText(self.template_file_url)
        else:
            warning = QMessageBox()
            warning.setWindowIcon(self.icon)
            warning.setWindowTitle('Ошибка')
            warning.setText('Выбран некорректный файл шаблона.')
            warning.setIcon(QMessageBox.Icon.Warning)
            warning.setStandardButtons(QMessageBox.StandardButton.Ok)
            info_button = warning.addButton('Подробнее', QMessageBox.ButtonRole.ActionRole)
            info_button.clicked.connect(self.show_template_info)
            warning.exec()

    def is_valid_pattern(self, template_url):
        """Файл проверки, что файл шаблона отчета является корректным"""
        try:
            if not (template_url.endswith('.docx') or template_url.endswith('.doc')):
                return False
            template = docx.Document(template_url)
            template_text = str()
            for paragraph in template.paragraphs:
                template_text += paragraph.text
            for temp_word in self.templates:
                if temp_word not in template_text:
                    return False
            return True
        except Exception as e:
            self.show_message(QMessageBox.Icon.Critical, 'Ошибка',
                              f'Произошли проблемы при работе с шаблоном.\nОшибка: {str(e)}')
        return False

    def create_report(self):
        """Создание отчета"""
        # Загрузка данных из заполненных полей
        self.templates['{university_name}'] = self.ui.universtyEdit.text()
        self.templates['{faculty_name}'] = self.ui.facultyEdit.text()
        if self.ui.workTypeBox.currentIndex() == 0:
            self.templates['{work_type}'] = 'лабораторной'
        elif self.ui.workTypeBox.currentIndex() == 1:
            self.templates['{work_type}'] = 'практической'
        self.templates['{work_num}'] = self.ui.workNumberBox.text()
        self.templates['{subject_name}'] = self.ui.subjectEdit.text()
        self.templates['{work_topic}'] = self.ui.topicEdit.text()
        self.templates['{group_num}'] = self.ui.groupEdit.text()
        self.templates['{student_name1}'] = '\n' + self.ui.studentEdit1.text()
        self.templates['{student_name2}'] = '-'
        self.templates['{student_name3}'] = '-'
        self.templates['{student_name4}'] = '-'
        self.templates['{student_name5}'] = '-'
        for i in range(0, self.students_count - 1):
            self.templates['{' + f'student_name{i + 2}' + '}'] = '\n' + self.students_edits[i].text()
        self.students_count = int(self.ui.studentCountLabel.text())
        self.templates['{teacher_post}'] = self.ui.postEdit.text()
        self.templates['{teacher_name}'] = self.ui.teacherEdit.text()
        self.templates['{city}'] = self.ui.cityBox.currentText()
        self.templates['{year}'] = str(datetime.now().year)
        self.templates['{purpose_work}'] = self.ui.purposeText.toPlainText()
        self.templates['{work_progress}'] = self.ui.progressText.toPlainText()
        self.templates['{conclusion}'] = self.ui.conclusionText.toPlainText()
        self.template_file_url = self.ui.templateFileEdit.text()
        self.report_file_name = self.ui.fileNameEdit.text()
        self.report_format = self.ui.formatBox.currentText()

        for template in (*self.templates.values(), self.template_file_url, self.report_file_name):
            if not template.strip():
                self.show_message(QMessageBox.Icon.Warning, 'Внимание', 'Заполните все обязательные поля.')
                return

        # Очистка полей имен студентов, которые не используются
        for i in range(5, self.students_count, -1):
            self.templates['{' + f'student_name{i}' + '}'] = ''

        report_extension = str()
        if 'docx' in self.report_format:
            report_extension = '.docx'
        elif 'doc' in self.report_format:
            report_extension = '.doc'
        elif 'pdf' in self.report_format:
            report_extension = '.pdf'

        try:
            report = docx.Document(self.template_file_url)

            # Замена полей в шаблоне
            for paragraph in report.paragraphs:
                text = paragraph.text
                # Добавление пространства перед городом для выравнивания титульного листа
                if '{city}' in text:
                    [paragraph.insert_paragraph_before('') for _ in range(5 - self.students_count)]
                # Замена по шаблону
                for template in self.templates:
                    if template in text:
                        text = text.replace(template, self.templates[template])
                        paragraph.text = text

            # Установка стилей
            style = report.styles['Normal']  # объект нормального стиля
            style.font.name = 'Times New Roman'  # шрифт
            style.font.size = Pt(14)  # размер шрифта

            # Сохранение сгенерированного отчета
            is_success_creation = False
            if report_extension == '.pdf':
                # Сохранение в формате PDF
                report_save_path = QFileDialog.getSaveFileName(self, 'Сохранение отчета',
                                                               f"{os.environ['USERPROFILE']}\\Documents\\" +
                                                               self.report_file_name + '.docx')[0]
                if report_save_path:
                    report.save(report_save_path)
                    docx2pdf.convert(report_save_path, report_save_path.replace('.docx', '.pdf'))
                    os.remove(report_save_path)
                    is_success_creation = True
            else:
                # Сохранение в формате Word
                report_save_path = QFileDialog.getSaveFileName(self, 'Сохранение отчета',
                                                               f"{os.environ['USERPROFILE']}\\Documents\\" +
                                                               self.report_file_name + report_extension)[0]
                if report_save_path:
                    report.save(report_save_path)
                    is_success_creation = True

            if is_success_creation:
                # Сохранение данных после успешного создания отчета
                self.save_data()
                self.show_message(QMessageBox.Icon.NoIcon, 'Готово', 'Отчет успешно сгенерирован!')

        except Exception as e:
            self.show_message(QMessageBox.Icon.Critical, 'Ошибка',
                              f'Произошли проблемы при генерации отчета.\nОшибка: {str(e)}')

    def show_message(self, icon, title, text):
        """Отображение всплывающих окон"""
        message = QMessageBox()
        message.setWindowIcon(self.icon)
        message.setIcon(icon)
        message.setWindowTitle(title)
        message.setText(text)
        message.exec()

    def show_template_info(self):
        """Отображение окна с информацией о шаблоне отчета"""
        info = QMessageBox()
        info.setWindowIcon(self.icon)
        info.setWindowTitle('Инфо')
        info.setIcon(QMessageBox.Icon.Information)
        info.setText('Файл-шаблон для генерации отчетов должен иметь формат документа Word и следующие поля:')
        info.setInformativeText("""
        {university_name} - полное наименование учебного заведения
        {faculty_name} - название кафедры
        {work_type} - вид работы (лабораторная, практическая)
        {work_num} - номер работы
        {subject_name} - название дисциплины
        {work_topic} - тема работы
        {group_num} - номер группы студента
        {student_name1} - Фамилия И.О. студента 1
        {student_name2} - Фамилия И.О. студента 2
        {student_name3} - Фамилия И.О. студента 3 
        {student_name4} - Фамилия И.О. студента 4
        {student_name5} - Фамилия И.О. студента 5 
        {teacher_post} - должность преподавателя
        {teacher_name} - Фамилия И.О. преподавателя
        {city} - город
        {year} - текущий год
        {purpose_work} - содержание цели работы
        {work_progress} - содержание хода работы
        {conclusion} - содержание вывода
        """)
        info.exec()

    @staticmethod
    def resource_path(relative_path):
        """Получает абсолютный путь к ресурсу"""
        if getattr(sys, 'frozen', False):
            # Если приложение собрано в .exe
            base_path = sys._MEIPASS
        else:
            # Усли скрипт запущен напрямую
            base_path = os.path.abspath(".")

        return os.path.join(base_path, relative_path)